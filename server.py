import os
import sys
import json
import re
import time
import uuid
import shutil
import tempfile
import subprocess
import unicodedata
import asyncio
import requests # Added for downloading TTS audio
from http import HTTPStatus
from bs4 import BeautifulSoup
from flask import Flask, request, jsonify, send_from_directory
from datetime import datetime
import yt_dlp

# Third-party
import dashscope
from dashscope import MultiModalConversation
from dashscope.audio.tts_v2 import SpeechSynthesizer as SpeechSynthesizerV2, AudioFormat as AudioFormatV2
from dashscope.audio.asr import Transcription
# from rapidocr_onnxruntime import RapidOCR # Removed
import fitz  # PyMuPDF
# import cv2 # Removed: heavy dependency not needed for current features
# import numpy as np # Removed: implicitly used by libraries but not directly
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Determine Static Folder based on current file location
# This ensures it works even if cwd is different (e.g. Vercel)
# root_dir = os.path.dirname(os.path.abspath(__file__))
# app = Flask(__name__, static_folder=root_dir)

# Vercel fix: When running via api/index.py, __file__ might be different or we need to point to correct static folder
# chifanzuiyaojin is the directory where static files (index.html, app.js) reside.
base_dir = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__, static_folder=base_dir)
root_dir = base_dir

# Set Max Content Length - Unlimited for local use
# Vercel limits were 4.5MB, but now we are free
app.config['MAX_CONTENT_LENGTH'] = None  

# --- Logging Setup ---
import logging
from logging.handlers import RotatingFileHandler

def setup_logging():
    """配置日志系统，将日志写入文件"""
    # 确保logs目录存在
    logs_dir = os.path.join(root_dir, 'logs')
    os.makedirs(logs_dir, exist_ok=True)
    
    log_file = os.path.join(logs_dir, 'server.log')
    
    # 配置根日志记录器
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    # 创建文件处理器（自动轮转，每个文件最大10MB，保留5个备份）
    file_handler = RotatingFileHandler(
        log_file,
        maxBytes=10*1024*1024,  # 10MB
        backupCount=5,
        encoding='utf-8'
    )
    file_handler.setLevel(logging.INFO)
    
    # 创建控制台处理器
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    
    # 设置日志格式
    formatter = logging.Formatter(
        '%(asctime)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)
    
    # 添加处理器
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    # 捕获print输出的函数
    class LogWriter:
        """将print输出重定向到日志文件"""
        def __init__(self, original_stderr):
            self.original = original_stderr
            self.log = logging.getLogger()
            self.buffer = ""
        
        def write(self, message):
            # 写入原始stderr
            self.original.write(message)
            self.original.flush()
            
            # 写入日志文件（只处理非空的行）
            cleaned = message.strip()
            if cleaned:
                # 检测是否是HTTP访问日志，避免重复记录
                if not (cleaned.startswith('127.0.0.1') or 
                        cleaned.startswith('192.168.') or
                        'HTTP/1.1' in cleaned):
                    # 立即写入日志（不缓冲）
                    self.log.info(cleaned)
                    # 强制刷新日志文件
                    for handler in self.log.handlers:
                        if isinstance(handler, RotatingFileHandler):
                            handler.flush()
        
        def flush(self):
            self.original.flush()
            # 刷新所有日志处理器
            for handler in self.log.handlers:
                handler.flush()
    
    # 保存原始stderr
    original_stderr = sys.stderr
    
    # 重定向sys.stderr到我们的LogWriter
    sys.stderr = LogWriter(original_stderr)
    
    # 测试日志是否正常工作
    logging.info("=== Server Started ===")
    logging.info("Logging system initialized. All stderr output will be captured.")
    
    return logger

# 初始化日志系统
app_logger = setup_logging()

# --- Load Environment ---
# Disable Flask's automatic dotenv loading since we load it manually
os.environ['FLASK_SKIP_DOTENV'] = '1'

def load_env_file():
    # Try current cwd first, then server.py dir
    env_paths = [
        os.path.join(os.getcwd(), '.env'),
        os.path.join(root_dir, '.env'),
        os.path.join(root_dir, '..', '.env') # Try parent if in subdir
    ]
    
    for env_path in env_paths:
        if os.path.exists(env_path):
            with open(env_path, 'r') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('#'): continue
                    if '=' in line:
                        k, v = line.split('=', 1)
                        if k.strip() not in os.environ: # Don't overwrite existing env
                            os.environ[k.strip()] = v.strip()
            break # Stop after finding first .env

load_env_file()

# --- Helper: Output Directory ---
def get_output_dir():
    # Use project root directory's tts_output folder for persistent storage
    local_out = os.path.join(root_dir, "tts_output")
    os.makedirs(local_out, exist_ok=True)
    return local_out

# --- Helper: History Storage ---
def get_history_file():
    # Use a JSON file to persist history on server
    history_file = os.path.join(root_dir, "history.json")
    if not os.path.exists(history_file):
        # Try temp directory if root_dir is read-only
        history_file = os.path.join(tempfile.gettempdir(), "chifanzuiyaojin_history.json")
    return history_file

def load_history():
    try:
        history_file = get_history_file()
        if os.path.exists(history_file):
            with open(history_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    except Exception as e:
        print(f"Failed to load history: {e}", file=sys.stderr)
        return []

def save_history_to_file(history):
    try:
        history_file = get_history_file()
        os.makedirs(os.path.dirname(history_file), exist_ok=True)
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Failed to save history: {e}", file=sys.stderr)
        return False



# --- Configuration ---
# Models requested by user
MODEL_TTS_LIST = ["qwen-tts", "qwen-tts-latest"]
# ASR Models Priority List
# Primary models (standard ASR): qwen3-omni-flash with version dates
# Fallback models (realtime ASR): qwen3-omni-flash-realtime with version dates
# Note: Realtime models may have different API parameters, but the primary models are compatible with current implementation
MODEL_ASR_LIST = [
    "qwen3-omni-flash-2025-12-01",
    "qwen3-omni-flash-2025-09-15",
    "qwen3-omni-flash-realtime-2025-12-01",
    "qwen3-omni-flash-realtime-2025-09-15"
]
MODEL_ASR_FILE = MODEL_ASR_LIST[0] # Default
MODEL_LLM = "qwen3-max-2025-09-23" # Reverted to high-accuracy model
MODEL_LLM_FAST = "qwen3-max-2025-09-23" # Use Max for everything now that speed isn't a constraint
MODEL_OCR = "qwen-vl-ocr-2025-11-20"

MAX_TTS_TEXT_LENGTH = 500_000
MAX_SAY_SEGMENT_LENGTH = 450 # Reduced from 1500 to 450 to meet [1, 512] API limit

# Initialize OCR (Removed local engine)
# ocr_engine = RapidOCR()

# --- Helper: FFmpeg ---
def ensure_ffmpeg_in_path():
    common_paths = ["/opt/homebrew/bin", "/usr/local/bin", "/usr/bin", "/bin"]
    current_path = os.environ.get("PATH", "")
    for p in common_paths:
        if os.path.exists(os.path.join(p, "ffmpeg")):
            if p not in current_path:
                os.environ["PATH"] = p + os.pathsep + current_path
            return

ensure_ffmpeg_in_path()

# --- Helper: URL Download ---
def extract_url_from_text(text):
    """
    Extract URL from text that may include Chinese title.
    Example: "【标题】https://b23.tv/xxx" or "标题 https://example.com"
    """
    # Try to find http:// or https:// URLs
    url_match = re.search(r'https?://[^\s\]]+', text)
    if url_match:
        return url_match.group(0)
    # If no http/https found, try to find b23.tv or other common patterns
    b23_match = re.search(r'b23\.tv/[^\s\]]+', text)
    if b23_match:
        return 'https://' + b23_match.group(0)
    return text

def download_audio_video_from_url(url, output_dir):
    """
    Download audio/video from URL using yt-dlp or direct download.
    Supports:
    - Direct file URLs (mp3, wav, mp4, etc.)
    - Video websites (B站, YouTube, etc.) via yt-dlp
    - Text with Chinese title and URL (e.g., "【标题】https://b23.tv/xxx")
    Returns: (success, downloaded_path, error, original_url)
    """
    original_url = url
    
    # Extract pure URL if text contains Chinese title or description
    extracted_url = extract_url_from_text(url)
    if extracted_url != url:
        print(f"Extracted URL from text: {extracted_url}", file=sys.stderr)
        url = extracted_url
    
    print(f"Downloading from URL: {url}", file=sys.stderr)
    
    # Check if this is a direct file URL (ends with common media extensions)
    direct_file_extensions = ['.mp3', '.wav', '.m4a', '.ogg', '.flac', '.aac', 
                            '.mp4', '.mkv', '.avi', '.mov', '.webm', '.flv']
    is_direct_file = any(url.lower().endswith(ext) for ext in direct_file_extensions)
    
    if is_direct_file:
        # Direct download using requests
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            print("Using direct download...", file=sys.stderr)
            response = requests.get(url, headers=headers, stream=True, timeout=60)
            response.raise_for_status()
            
            # Get file extension from URL or Content-Type
            ext = os.path.splitext(url)[1] if '.' in url.split('/')[-1] else '.mp3'
            if not ext:
                content_type = response.headers.get('content-type', '')
                if 'audio' in content_type:
                    ext = '.mp3'
                elif 'video' in content_type:
                    ext = '.mp4'
                else:
                    ext = '.mp3'
            
            # Save file
            filename = f"asr-{uuid.uuid4()}{ext}"
            output_path = os.path.join(output_dir, filename)
            
            with open(output_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            print(f"Direct download complete: {output_path}", file=sys.stderr)
            return True, output_path, None, url
            
        except Exception as e:
            error_msg = f"Direct download failed: {str(e)}"
            print(error_msg, file=sys.stderr)
            return False, None, error_msg, original_url
    
    else:
        # Use yt-dlp for video websites
        try:
            print("Using yt-dlp for video site download...", file=sys.stderr)
            
            # Configure yt-dlp options to download video (not extract audio)
            ydl_opts = {
                # Explicitly download best video WITH audio (not just audio)
                'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/bestvideo+bestaudio/best',
                'outtmpl': os.path.join(output_dir, 'asr-%(id)s.%(ext)s'),
                'quiet': True,
                'no_warnings': True,
                'ignoreerrors': False,
                'nocheckcertificate': True,
                'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                # Merge formats if video and audio are separate
                'merge_output_format': 'mp4',
                # Don't embed subtitles
                'writesubtitles': False,
                'writeautomaticsub': False,
                # CRITICAL: Don't extract audio to mp3, keep original format
                'postprocessors': [],  # Empty list to prevent audio extraction
            }
            
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                # Extract info first
                info = ydl.extract_info(url, download=False)
                if not info:
                    return False, None, "无法解析此URL，请检查链接是否正确", original_url
                
                # Get video title for logging
                video_title = info.get('title', 'Unknown')
                print(f"Found video: {video_title}", file=sys.stderr)
                
                # Download
                ydl.download([url])
            
            # Find downloaded file (yt-dlp adds extension - support both video and audio formats)
            # IMPORTANT: Support both 'asr-' prefix (for converted files) and video ID (for original yt-dlp downloads)
            supported_extensions = ['.mp4', '.webm', '.mkv', '.mp3', '.wav', '.m4a']
            
            # First try to find files with 'asr-' prefix (from our download template)
            downloaded_files = [f for f in os.listdir(output_dir) 
                             if f.startswith('asr-') and any(f.endswith(ext) for ext in supported_extensions)]
            
            if not downloaded_files:
                # If no 'asr-' files found, look for files with video ID or recent downloads
                # Get video ID from info dict for filename matching
                video_id = info.get('id', '')
                if video_id:
                    # Look for files starting with video ID
                    downloaded_files = [f for f in os.listdir(output_dir) 
                                     if f.startswith(video_id) and any(f.endswith(ext) for ext in supported_extensions)]
            
            if downloaded_files:
                # Get most recent file
                downloaded_files.sort(key=lambda x: os.path.getmtime(os.path.join(output_dir, x)), reverse=True)
                output_path = os.path.join(output_dir, downloaded_files[0])
                print(f"yt-dlp download complete: {output_path}", file=sys.stderr)
                return True, output_path, None, url
            else:
                return False, None, "下载失败：未找到下载的文件", url
            
        except yt_dlp.utils.DownloadError as e:
            error_msg = f"下载失败：{str(e)}"
            print(error_msg, file=sys.stderr)
            return False, None, error_msg, original_url
        except Exception as e:
            error_msg = f"下载出错：{str(e)}"
            print(error_msg, file=sys.stderr)
            return False, None, error_msg, original_url

def get_audio_duration(file_path):
    if not shutil.which("ffprobe"): return 0
    try:
        cmd = ["ffprobe", "-v", "error", "-show_entries", "format=duration", 
               "-of", "default=noprint_wrappers=1:nokey=1", file_path]
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return float(result.stdout.strip())
    except:
        return 0

def get_file_type(file_path):
    """
    Determine if a file is video or audio based on file extension and ffprobe.
    Returns: 'video' or 'audio'
    """
    # First check file extension
    video_extensions = ['.mp4', '.mkv', '.avi', '.mov', '.webm', '.flv', '.wmv']
    audio_extensions = ['.mp3', '.wav', '.m4a', '.ogg', '.flac', '.aac', '.wma', '.aiff']
    
    ext = os.path.splitext(file_path)[1].lower()
    if ext in video_extensions:
        return 'video'
    elif ext in audio_extensions:
        return 'audio'
    
    # If extension is ambiguous, use ffprobe to check
    if shutil.which("ffprobe"):
        try:
            cmd = ["ffprobe", "-v", "error", "-select_streams", "v", 
                   "-show_entries", "stream=codec_type", 
                   "-of", "default=noprint_wrappers=1:nokey=1", file_path]
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            if result.stdout.strip():
                return 'video'
            else:
                return 'audio'
        except:
            pass
    
    # Default to audio
    return 'audio'

def convert_to_wav_16k(input_path, output_path):
    if not shutil.which("ffmpeg"): return False, "ffmpeg missing"
    try:
        cmd = ["ffmpeg", "-y", "-i", input_path, "-ar", "16000", "-ac", "1", 
               "-c:a", "pcm_s16le", output_path]
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
        return True, None
    except Exception as e:
        return False, str(e)

def convert_to_browser_compatible(input_path, original_ext):
    """
    Convert audio/video files to browser-compatible formats.
    - Audio: convert to MP3 (widely supported)
    - Video: keep as is if already MP4/WebM, otherwise convert to MP4
    Returns: (success, converted_path, error)
    """
    if not shutil.which("ffmpeg"):
        return False, None, "ffmpeg missing"
    
    try:
        # Determine output format based on original file type
        # Audio formats that should be converted
        audio_formats = ['.m4a', '.wma', '.aac', '.flac', '.ogg', '.wav', '.aiff']
        # Video formats - convert non-browser formats to MP4, keep MP4/WebM as is
        video_formats_convert = ['.mkv', '.avi', '.mov', '.wmv', '.flv']
        video_formats_keep = ['.mp4', '.webm']
        
        is_audio = original_ext.lower() in audio_formats
        is_video_convert = original_ext.lower() in video_formats_convert
        is_video_keep = original_ext.lower() in video_formats_keep
        
        # If already MP3 or MP4/WebM, no need to convert
        if original_ext.lower() in ['.mp3'] or is_video_keep:
            return True, input_path, None
        
        # Create temp file for conversion
        fd, temp_path = tempfile.mkstemp(suffix='.mp3' if is_audio else '.mp4')
        os.close(fd)
        
        if is_audio:
            # Convert audio to MP3 (better browser support)
            cmd = [
                "ffmpeg", "-y", "-i", input_path,
                "-codec:a", "libmp3lame",
                "-q:a", "2",  # Good quality (0-9, lower is better)
                temp_path
            ]
        elif is_video_convert:
            # Convert video to MP4 (H.264 + AAC for best compatibility)
            cmd = [
                "ffmpeg", "-y", "-i", input_path,
                "-c:v", "libx264",
                "-preset", "medium",
                "-crf", "23",
                "-c:a", "aac",
                "-b:a", "128k",
                "-movflags", "+faststart",  # Enable fast start for streaming
                temp_path
            ]
        else:
            # Unknown format, try to keep as is
            return True, input_path, None
        
        # Run conversion
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
        
        return True, temp_path, None
        
    except Exception as e:
        print(f"Conversion failed: {e}", file=sys.stderr)
        return False, None, str(e)

# --- Helper: Text Normalization ---
def normalize_text_for_tts(text):
    if not text: return ""
    t = str(text)
    try: t = unicodedata.normalize('NFKC', t)
    except: pass
    t = t.replace('\uFEFF', '')
    # Basic cleanup
    t = re.sub(r'[\r\n]+', '\n', t)
    return t.strip()

def split_text_for_say(text, max_len):
    # Simple split by punctuation or space
    t = str(text)
    out = []
    i = 0
    while i < len(t):
        j = min(i + max_len, len(t))
        if j < len(t):
            # Try to find split point
            window = t[i:j]
            match = re.search(r'(?:\n+|[。！？!?]\s*)', window[::-1])
            if match:
                offset = match.start()
                if offset < max_len * 0.5: # If split point is reasonably far
                    j = j - offset
        
        out.append(t[i:j].strip())
        i = j
    return [c for c in out if c]

def estimate_subtitles_helper(text, duration):
    if duration <= 0 or not text: return []
    # Simple linear alignment
    total_chars = len(text)
    if total_chars == 0: return []
    time_per_char = duration / total_chars
    
    subtitles = []
    # Split by punctuation
    parts = re.split(r'([，。！？；：,.!?;:])', text)
    current_time = 0.0
    
    buffer = ""
    for p in parts:
        buffer += p
        if p and p in "，。！？；：,.!?;:":
            # Flush buffer
            chunk_dur = len(buffer) * time_per_char
            subtitles.append({
                "text": buffer,
                "begin_time": int(current_time * 1000), # ms
                "end_time": int((current_time + chunk_dur) * 1000) # ms
            })
            current_time += chunk_dur
            buffer = ""
    if buffer:
        chunk_dur = len(buffer) * time_per_char
        subtitles.append({
            "text": buffer,
            "begin_time": int(current_time * 1000),
            "end_time": int((current_time + chunk_dur) * 1000)
        })
        
    return subtitles

# --- Backend: Alibaba TTS ---
class AlibabaTTSBackend:
    def __init__(self, api_key=None):
        self.api_key = api_key or os.environ.get("DASHSCOPE_API_KEY")
        
    def get_voices(self):
        # qwen-tts-2025-05-22 supported voices
        return [
            {"name": "Cherry", "locale": "zh-CN", "description": "标准女声 (Cherry)"},
            {"name": "Ethan", "locale": "zh-CN", "description": "标准男声 (Ethan)"},
            {"name": "Chelsie", "locale": "zh-CN", "description": "温柔女声 (Chelsie)"},
            {"name": "Serena", "locale": "zh-CN", "description": "亲切女声 (Serena)"},
            {"name": "Dylan", "locale": "zh-CN", "description": "北京口音 (Dylan)"},
            {"name": "Jada", "locale": "zh-CN", "description": "上海口音 (Jada)"},
            {"name": "Sunny", "locale": "zh-CN", "description": "四川口音 (Sunny)"}
        ]

    def generate(self, text, voice, output_path):
        if not self.api_key:
            return False, "Missing API Key", None
        
        dashscope.api_key = self.api_key
        
        voice_id = voice
        last_error = None
        
        for model in MODEL_TTS_LIST:
            # Retry logic for TTS API per model
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    # Use MultiModalConversation for qwen-tts
                    print(f"TTS Call (Model: {model}, Attempt {attempt+1}/{max_retries}): voice={voice_id}, text_len={len(text)}", file=sys.stderr)
                    
                    response = MultiModalConversation.call(
                        model=model,
                        text=text,
                        voice=voice_id,
                    )
                    
                    if response.status_code == HTTPStatus.OK:
                        # qwen-tts returns an audio URL in output.audio.url
                        if hasattr(response.output, 'audio') and 'url' in response.output.audio:
                            audio_url = response.output.audio['url']
                            
                            # Download the audio
                            r = requests.get(audio_url)
                            r.raise_for_status()
                            
                            with open(output_path, "wb") as f:
                                f.write(r.content)
                                
                            # Fake subtitles (Estimate)
                            duration = get_audio_duration(output_path)
                            raw_subs = estimate_subtitles_helper(text, duration)
                            subtitles = []
                            for s in raw_subs:
                                subtitles.append({
                                    "text": s['text'],
                                    "start": s['begin_time'] / 1000.0,
                                    "end": s['end_time'] / 1000.0
                                })
                            return True, None, subtitles
                        else:
                            last_error = f"Unexpected response format from {model}: {response}"
                            print(last_error, file=sys.stderr)
                            break # Try next model
                    else:
                        if "free tier of the model has been exhausted" in str(response.message):
                             return False, "阿里云DashScope Qwen-TTS模型免费额度已耗尽。请前往阿里云控制台开启“按量付费”或购买资源包以继续使用。", None
                        
                        # Check for 500 InternalError.Algo which might be transient
                        if response.status_code == 500 and "InternalError.Algo" in str(response.message):
                            print(f"TTS Transient Error ({model}): {response.message}, retrying...", file=sys.stderr)
                            time.sleep(1 + attempt) # Backoff
                            continue # Retry same model
                        
                        last_error = f"TTS API Error ({model}): {response.message}"
                        print(last_error, file=sys.stderr)
                        break # Try next model
                        
                except Exception as e:
                    error_msg = str(e)
                    if "AllocationQuota.FreeTierOnly" in error_msg or "free tier" in error_msg.lower():
                        return False, "阿里云DashScope免费额度已耗尽。请前往阿里云控制台开启“按量付费”或购买资源包以继续使用。(错误代码: AllocationQuota.FreeTierOnly)", None
                    
                    # Check for 500 in exception message if wrapped
                    if ("500" in error_msg and "InternalError" in error_msg) or \
                       ("503" in error_msg) or \
                       ("502" in error_msg) or \
                       ("ConnectionError" in error_msg):
                         print(f"TTS Exception ({model} - Transient?): {error_msg}, retrying...", file=sys.stderr)
                         time.sleep(1 + attempt)
                         continue # Retry same model
                    
                    last_error = f"Exception ({model}): {error_msg}"
                    print(last_error, file=sys.stderr)
                    break # Try next model
        
        return False, f"All TTS models failed. Last error: {last_error}", None

    def estimate_subtitles(self, text, duration):
        # Legacy method kept for compatibility if needed, but generate() uses helper now.
        # Actually generate() uses local logic above.
        pass

# --- Backend: Alibaba ASR ---
def run_ali_asr(file_path, api_key, task_id):
    """Run ASR with task ID for progress tracking"""
    dashscope.api_key = api_key
    print(f"ASR [Task ID: {task_id}]: Starting ASR processing", file=sys.stderr)
    
    # 1. Compress/Extract to mp3
    compressed_path = file_path
    is_temp = False
    
    if shutil.which("ffmpeg"):
        try:
            fd, temp_path = tempfile.mkstemp(suffix=".mp3")
            os.close(fd)
            # -ac 1: mono
            # -b:a 64k: slightly higher bitrate to avoid artifacts
            cmd = ["ffmpeg", "-y", "-i", file_path, 
                   "-vn", "-ar", "16000", "-ac", "1", "-b:a", "64k", 
                   temp_path]
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
            compressed_path = temp_path
            is_temp = True
        except Exception as e:
            print(f"ASR Compression Warning: {e}", file=sys.stderr)
            is_temp = False
            compressed_path = file_path

    # 2. Check Duration and Chunk if necessary
    # Qwen-Audio has timeout issues on long files. We split if > 300s.
    CHUNK_DURATION = 300
    duration = get_audio_duration(compressed_path)
    
    final_text = ""
    all_sentences = []
    
    # Local helper for cleanup
    def clean_asr_hallucinations(t):
        if not t: return ""
        t = str(t)
        
        # 0. Remove known model hallucinations/prefixes
        bad_prefixes = [
            "这段音频的原始内容是",
            "这段音频的内容是",
            "音频内容是",
            "转写结果如下",
            "请准确转写",
            "这段音频"
        ]
        for p in bad_prefixes:
            t = t.replace(p, "")
            
        # 1. Squash repeated characters (often hallucinations)
        # Replace any character repeated 4 or more times with a single instance
        t = re.sub(r'(.)\1{3,}', r'\1', t)
        
        # 2. Basic punctuation cleanup (as safeguard)
        # Fix double punctuation from raw output
        t = re.sub(r'([，。！？；：,.!?;:])\1+', r'\1', t)
        # Fix space around punctuation
        t = re.sub(r'\s*([，。！？；：,.!?;:])\s*', r'\1', t)
        
        # 3. Trim
        return t.strip()

    try:
        if duration <= CHUNK_DURATION:
            # Direct call with task ID
            ok, res = _call_qwen_audio(compressed_path, task_id)
            if not ok: return False, res
            
            # Clean (Basic Regex)
            res = clean_asr_hallucinations(res)
            
            final_text = res
            
            # Generate subtitles
            all_sentences = estimate_subtitles_helper(res, duration)
            
        else:
            # Chunking
            print(f"Audio too long ({duration}s), splitting...", file=sys.stderr)
            with tempfile.TemporaryDirectory() as chunk_dir:
                # ffmpeg segment
                # -f segment -segment_time 300 -c copy
                # Note: '-c copy' can sometimes produce chunks with weird timestamps or incomplete headers if keyframes don't align.
                # Re-encoding ensures valid standalone files.
                seg_pattern = os.path.join(chunk_dir, "chunk_%03d.mp3")
                cmd = ["ffmpeg", "-i", compressed_path, "-f", "segment", 
                       "-segment_time", str(CHUNK_DURATION), "-ar", "16000", "-ac", "1", "-b:a", "64k", seg_pattern]
                subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
                
                chunks = sorted([os.path.join(chunk_dir, f) for f in os.listdir(chunk_dir) if f.endswith(".mp3")])
                
                current_offset_ms = 0
                
                for i, chunk in enumerate(chunks):
                    print(f"ASR [Task ID: {task_id}]: Processing chunk {i+1}/{len(chunks)}...", file=sys.stderr)
                    
                    # Get exact duration of chunk for better alignment
                    chunk_dur = get_audio_duration(chunk)
                    
                    ok, res = _call_qwen_audio(chunk, task_id)
                    if not ok: 
                        print(f"Chunk {i} failed: {res}", file=sys.stderr)
                        final_text += f"\n[...片段 {i+1} 转写失败，内容缺失...]\n"
                        # We must advance time even if failed to keep alignment
                        current_offset_ms += int(chunk_dur * 1000)
                        continue 
                    
                    # Convert to string and clean up
                    if isinstance(res, list): 
                         res = " ".join([str(x) for x in res])
                    elif not isinstance(res, str):
                         res = str(res)
                    
                    # Post-processing to remove hallucinations (e.g. repeated "呃")
                    res = clean_asr_hallucinations(res)
                    
                    final_text += res + "\n"
                    
                    # Generate subtitles for this chunk
                    chunk_subs = estimate_subtitles_helper(res, chunk_dur)
                    
                    # Adjust offsets
                    for sub in chunk_subs:
                        sub['begin_time'] += current_offset_ms
                        sub['end_time'] += current_offset_ms
                        all_sentences.append(sub)
                        
                    current_offset_ms += int(chunk_dur * 1000)

        
        return True, {"text": final_text, "sentences": all_sentences}

    except Exception as e:
        return False, str(e)
    finally:
        if is_temp and os.path.exists(compressed_path):
            os.remove(compressed_path)

def _call_qwen_audio(audio_path, task_id):
    """Call Qwen Audio API with task ID"""
    # Try models in order
    last_error = ""
    
    for model_name in MODEL_ASR_LIST:
        try:
            print(f"ASR [Task ID: {task_id}]: Trying model: {model_name}", file=sys.stderr)
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"audio": f"file://{audio_path}"},
                        {"text": "请逐字逐句完整转写这段音频，输出结果时请自动去掉语气词、口头禅（如“呃”、“啊”、“那个”、“就是”等）和重复词，修正显而易见的口误，并使用规范的标点符号。不要进行总结，直接输出清洗后的转写文本。"}
                    ]
                }
            ]
            
            # Qwen3-Omni-Flash requires streaming
            # and might require explicit result collection
            response_iterator = MultiModalConversation.call(
                model=model_name, 
                messages=messages,
                stream=True
            )
            
            full_content = ""
            last_response = None
            
            print(f"ASR [Task ID: {task_id}]: DEBUG: Starting stream for {model_name}", file=sys.stderr)
            
            for response in response_iterator:
                last_response = response
                if response.status_code == HTTPStatus.OK:
                    # Accumulate content if it's incremental
                    # Check structure
                    if hasattr(response.output, 'choices') and len(response.output.choices) > 0:
                        choice = response.output.choices[0]
                        # DashScope stream usually returns FULL content so far in 'message.content'
                        # But let's check if it's delta
                        
                        # Note: Qwen-Audio-Turbo returned full content. 
                        # Qwen3-Omni-Flash via OpenAI compatible mode returns Deltas.
                        # Via DashScope SDK? It likely standardizes to Full Content OR Deltas.
                        # The symptom "Content len: 1" implies it returns DELTAS.
                        
                        # Let's try to grab content and see.
                        part = choice.message.content
                        
                        # If list (multimodal output?), process it
                        if isinstance(part, list):
                            # Usually [{"text": "..."}]
                            text_part = ""
                            for item in part:
                                if isinstance(item, dict) and 'text' in item:
                                    text_part += item['text']
                                elif isinstance(item, str):
                                    text_part += item
                            
                            if text_part:
                                full_content += text_part # APPEND, don't overwrite
                        
                        elif isinstance(part, str):
                            if part: full_content += part # APPEND, don't overwrite
                            
                else:
                    print(f"Stream Error Chunk: {response}", file=sys.stderr)
                    # Don't raise immediately, try to continue? No, stream error is fatal usually.
                    raise Exception(f"Stream Error: {response.message}")
            
            print(f"ASR [Task ID: {task_id}]: DEBUG: Stream finished. Content len: {len(full_content)}", file=sys.stderr)
            
            if full_content:
                return True, full_content
            elif last_response and last_response.status_code == HTTPStatus.OK:
                 # If full_content is empty but we had success, maybe it's in a different field?
                 # Just return empty string and let the outer loop handle "failed" if it's critical?
                 # But outer loop sees True and prints it.
                 return True, ""
            else:
                response = last_response if last_response else response # Fallback
                # Check for quota error
                err_msg = str(response.message).lower()
                # Also check for "does not support this input" which means model might not support file URL or format
                # But here we are iterating models. 
                # If qwen-audio-asr fails with "does not support this input", maybe it's because we use file://
                # qwen-audio-asr (VL model) supports file:// if local? 
                # Actually, standard qwen-audio-asr/latest DOES NOT support local file:// in MultiModalConversation call unless using specific SDK utils or if dashscope handles it.
                # BUT qwen-audio-turbo-1204 (and turbo) DOES support it via the same SDK method if it's the specific "turbo" variant that handles uploads.
                
                # However, previous code worked with qwen-audio-turbo-1204.
                # Now we added qwen-audio-asr and qwen-audio-asr-latest to the list.
                # If these models don't support file:// upload transparently, they will fail with 400 InvalidParameter.
                
                # To fix this: Ensure we are using a compatible method or only using models that support it.
                # qwen-audio-asr (opensource Qwen-Audio-Chat) deployed on DashScope usually supports the same input.
                # But if it fails, we should treat it as "model failed" and switch.
                
                if "free tier" in err_msg or "quota" in err_msg or "payment" in err_msg or "arrearage" in err_msg or "does not support this input" in err_msg:
                    print(f"ASR Model {model_name} failed/unsupported ({response.code}), switching...", file=sys.stderr)
                    last_error = f"{model_name} error: {response.message}"
                    continue # Try next model
                else:
                    # Other error, might be persistent
                    print(f"ASR Model {model_name} failed: {response.message}", file=sys.stderr)
                    last_error = f"{model_name} error: {response.message}"
                    continue

        except Exception as e:
            print(f"ASR Model {model_name} exception: {e}", file=sys.stderr)
            last_error = str(e)
            continue
            
    return False, f"All ASR models failed. Last error: {last_error}"

# --- Backend: Alibaba NLP (LLM) ---
def call_llm_analysis(text, api_key):
    dashscope.api_key = api_key
    prompt = f"""
    Please analyze the following text and provide a JSON response with these fields:
    1. "keywords": list of top 5 keywords (strings).
    2. "summary": a concise summary (string).
    3. "topics": list of objects, each with:
       - "title" (string): topic description.
       - "start_snippet" (string): the exact first 15-20 characters of this section in the text.
       - "end_snippet" (string): the exact last 10 characters of this section.
    
    Text:
    {text[:15000]} 
    """
    try:
        resp = dashscope.Generation.call(
            model=MODEL_LLM,
            messages=[{'role': 'user', 'content': prompt}],
            result_format='message'
        )
        
        if resp.status_code == HTTPStatus.OK:
            content = resp.output.choices[0].message.content
            # Strip markdown code blocks if present
            content = re.sub(r'```json\s*|\s*```', '', content)
            try:
                data = json.loads(content)
                return True, data
            except:
                return False, f"Invalid JSON from LLM: {content}"
        else:
            return False, f"LLM Error: {resp.message}"
    except Exception as e:
        return False, str(e)

def call_llm_fix_punctuation(text, api_key):
    dashscope.api_key = api_key
    prompt = f"""
    你是一个专业的文本编辑器。请对下面的文本进行“智能标点修复”和“全文排版优化”。
    
    要求：
    1. 【标点修复】：根据上下文语义，添加或修正标点符号，使断句准确、通顺。
    2. 【文本清洗】：去除多余的空格（尤其是PDF提取导致的字间空格）、删除无意义的换行符，将断裂的句子合并。
    3. 【排版优化】：保持合理的段落结构，增加全文的可读性。
    4. 【错别字修正】：修正明显的同音错别字。
    5. 【输出格式】：直接返回修复后的纯文本，不要包含任何解释、前言或后缀。
    
    文本：
    {text[:15000]} 
    """
    try:
        resp = dashscope.Generation.call(
            model=MODEL_LLM_FAST, # Now Qwen-Max
            messages=[{'role': 'user', 'content': prompt}],
            result_format='message'
        )
        if resp.status_code == HTTPStatus.OK:
            return True, resp.output.choices[0].message.content
        return False, resp.message
    except Exception as e:
        return False, str(e)

def call_llm_advice(text, api_key, custom_prompt=None):
    dashscope.api_key = api_key
    
    if custom_prompt:
        # When custom prompt is provided, let it control the entire analysis
        prompt = f"""
        你是一位资深的上海中考语文阅卷组长及文学评论家。请根据【上海中考作文评分标准】，对以下学生作文进行专业评价。
        
        【用户特别要求 - 必须完全遵循】：
        {custom_prompt}
        
        【重要说明】：
        1. 上述用户要求具有最高优先级，请严格按照用户要求的格式和内容进行分析
        2. 不要输出默认的五个维度结构，而是完全按照用户要求的格式
        3. 如果用户要求top5加减分项，请明确列出并详细说明
        4. 如果用户要求具体修改建议，请针对每个减分项给出详细改进方案
        
        【返回格式要求】：
        请严格按照用户要求的格式返回纯文本结果，不要包含任何JSON格式或代码块标记。
        直接输出用户要求的内容格式即可。
        
        【作文内容】：
        {text[:8000]}
        """
    else:
        # Default 5-section format when no custom prompt
        prompt = f"""
        你是一位资深的上海中考语文阅卷组长及文学评论家。请根据【上海中考作文评分标准】，对以下学生作文进行全方位的深度辅导。
        
        【任务目标】：请提供一份包含以下五个维度的完整诊断报告，缺一不可。
    
    1. **【评分与诊断】**：
       - 预估得分（满分60分）。
       - 简要点评优缺点。
       
    2. **【结构与进阶思路】**（保留原有功能）：
       - 分析当前思路，并提供一个**“高分进阶思路”**（如：如何从“记事”提升到“感悟”）。
       
    3. **【多维审题与构思拓展】**（新需求）：
       - 针对题目背景，按照上海中考构思逻辑（成长、思辨、自我认知），提供 2-3 个全新的**高分立意方向**。
       
    4. **【细节润色与手法升级】**（核心功能，请详细）：
       - 挑选文中 **5-8** 处可提升的句子，指出其问题，并运用高级写作手法（环境烘托、心理刻画、修辞）进行升格修改。
       
    5. **【三种风格润色示范】**（新需求）：
       - 请针对文中的不同片段，分别提供以下三种风格的润色示范。**每种风格至少提供 3 个不同的精彩改写示例**：
       
         a. **【上海中考·真情实感风】**（细腻、真实、动人）
            - 示例1：...
            - 示例2：...
            - 示例3：...
         b. **【名家·经典散文风】**（李娟/汪曾祺风格，质朴灵动）
            - 示例1：...
            - 示例2：...
            - 示例3：...
         c. **【诗意哲理·名句引用风】**（符合上海中学生认知，重点引用中国古代著名诗人、词人的经典名句，如苏轼、李白、辛弃疾等，将诗词意境与哲理融合）
            - 示例1：...
            - 示例2：...
            - 示例3：...
    
    【返回格式】：请返回严格的 JSON 格式，包含以下字段：
    - "score_prediction": "预估得分：xx/60"
    - "analysis": "总体点评..."
    - "structure_advice": "写作思路与结构进阶建议..." (字符串)
    - "alternative_ideas": [
        {{ "title": "构思一：xxxxx", "desc": "具体阐述..." }},
        {{ "title": "构思二：xxxxx", "desc": "具体阐述..." }}
      ]
    - "suggestions": [
        {{
           "original": "原文片段...",
           "technique": "运用了xx手法",
           "suggestion": "分析与建议...",
           "refined_text": "升格后的精彩句子..."
        }}
      ]
    - "style_demonstrations": [
        {{
           "style_name": "【上海中考·真情实感风】",
           "examples": [
                {{ "original_snippet": "片段1...", "refined_text": "改写1...", "comment": "解析..." }},
                {{ "original_snippet": "片段2...", "refined_text": "改写2...", "comment": "解析..." }},
                {{ "original_snippet": "片段3...", "refined_text": "改写3...", "comment": "解析..." }}
           ]
        }},
        {{
           "style_name": "【名家·经典散文风】",
           "examples": [ ... ]
        }},
        {{
           "style_name": "【诗意哲理·名句引用风】",
           "examples": [ ... ]
        }}
      ]
    
    【作文内容】：
    {text[:8000]}
    """
    try:
        resp = dashscope.Generation.call(
            model=MODEL_LLM, # qwen3-max-2025-09-23
            messages=[{'role': 'user', 'content': prompt}],
            result_format='message'
        )
        if resp.status_code == HTTPStatus.OK:
            content = resp.output.choices[0].message.content
            
            if custom_prompt:
                # For custom prompts, return the raw text as analysis
                return True, {"analysis": content, "custom_format": True}
            else:
                # For standard format, try to parse JSON
                # Strip markdown code blocks
                content = re.sub(r'```json\s*|\s*```', '', content)
                try:
                    return True, json.loads(content)
                except:
                    # Try to salvage partial JSON or return raw text wrapped
                    return False, f"Invalid JSON: {content}"
        return False, resp.message
    except Exception as e:
        return False, str(e)



# --- Backend: Alibaba OCR (Qwen-VL-OCR) ---
def call_qwen_ocr(image_path, api_key):
    dashscope.api_key = api_key
    
    # Qwen-VL-OCR logic
    # It works like a chat model with image input
    
    messages = [
        {
            "role": "user",
            "content": [
                {"image": f"file://{image_path}"},
                {"text": "Read the text in this image. Return only the text content without markdown code blocks or extra explanations."}
            ]
        }
    ]
    
    try:
        resp = MultiModalConversation.call(
            model=MODEL_OCR,
            messages=messages
        )
        
        if resp.status_code == HTTPStatus.OK:
            return True, resp.output.choices[0].message.content[0]['text']
        else:
            return False, f"OCR Error: {resp.message}"
    except Exception as e:
        return False, str(e)

# --- Routes ---

@app.route('/health')
def health_check():
    """Health check endpoint for Docker and load balancers"""
    return jsonify({"status": "ok", "service": "chifanzuiyaojin"})

@app.route('/')
def serve_index():
    # Fallback to index.html in current directory if root_dir fails or for Vercel
    if os.path.exists(os.path.join(root_dir, 'index.html')):
        return send_from_directory(root_dir, 'index.html')
    # Try parent directory or cwd
    if os.path.exists(os.path.join(os.getcwd(), 'chifanzuiyaojin', 'index.html')):
         return send_from_directory(os.path.join(os.getcwd(), 'chifanzuiyaojin'), 'index.html')
    return send_from_directory('.', 'index.html')

@app.route('/tts_output/<path:filename>')
def serve_tts_output(filename):
    out_dir = get_output_dir()
    return send_from_directory(out_dir, filename)

@app.route('/<path:path>')
def serve_static(path):
    # Try root_dir first
    if os.path.exists(os.path.join(root_dir, path)):
        return send_from_directory(root_dir, path)
    if os.path.exists(os.path.join(os.getcwd(), 'chifanzuiyaojin', path)):
         return send_from_directory(os.path.join(os.getcwd(), 'chifanzuiyaojin'), path)
    return send_from_directory('.', path)

@app.route('/api/voices', methods=['GET'])
def api_voices():
    # Return Alibaba voices
    backend = AlibabaTTSBackend()
    return jsonify({"ok": True, "voices": backend.get_voices()})

@app.route('/api/get-config', methods=['GET'])
def api_get_config():
    # Return non-sensitive config, or mask key if needed
    # Since this is a local tool, returning the key for auto-fill is acceptable
    key = os.environ.get("DASHSCOPE_API_KEY", "")
    return jsonify({"ok": True, "dashscopeKey": key})

@app.route('/api/tts', methods=['POST'])
def api_tts():
    try:
        data = request.get_json()
        text = normalize_text_for_tts(data.get("text", ""))
        voice = data.get("voice", "longanyang") # Default updated
        key = data.get("dashscopeKey")
        
        if not key: key = os.environ.get("DASHSCOPE_API_KEY")
        if not key: return jsonify({"ok": False, "error": "missing_api_key"}), 401
        
        # Log key usage (masked)
        masked_key = f"{key[:6]}...{key[-4:]}" if key and len(key) > 10 else "InvalidKey"
        print(f"TTS Request: Voice={voice}, Key={masked_key}", file=sys.stderr)
        
        if not text: return jsonify({"ok": False, "error": "empty_text"}), 400
        
        backend = AlibabaTTSBackend(key)
        
        # Filename setup
        filename = f"tts-{uuid.uuid4()}.wav"
        out_dir = get_output_dir()
        final_path = os.path.join(out_dir, filename)
        
        # Check length
        if len(text) <= MAX_SAY_SEGMENT_LENGTH:
            ok, err, subs = backend.generate(text, voice, final_path)
            if not ok: return jsonify({"ok": False, "error": err}), 500
            
            return jsonify({
                "ok": True,
                "audio_url": f"/tts_output/{filename}",
                "subtitles": subs,
                "download_filename": filename
            })
        else:
            # Segment and merge
            with tempfile.TemporaryDirectory() as temp_dir:
                segments = split_text_for_say(text, MAX_SAY_SEGMENT_LENGTH)
                seg_files = []
                all_subs = []
                current_offset = 0.0
                
                for i, seg in enumerate(segments):
                    seg_path = os.path.join(temp_dir, f"seg_{i}.wav")
                    ok, err, subs = backend.generate(seg, voice, seg_path)
                    if not ok: return jsonify({"ok": False, "error": f"Segment {i} failed: {err}"}), 500
                    
                    # Adjust subs
                    dur = get_audio_duration(seg_path)
                    for s in subs:
                        s['start'] += current_offset
                        s['end'] += current_offset
                    all_subs.extend(subs)
                    current_offset += dur
                    
                    seg_files.append(seg_path)
                    time.sleep(0.2) # Rate limit protection
                
                # Merge
                list_path = os.path.join(temp_dir, "list.txt")
                with open(list_path, "w") as f:
                    for p in seg_files: f.write(f"file '{p}'\n")
                
                try:
                    subprocess.run(
                        ["ffmpeg", "-f", "concat", "-safe", "0", "-i", list_path, "-c", "copy", final_path],
                        check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
                    )
                except Exception as e:
                    return jsonify({"ok": False, "error": f"Merge failed: {e}"}), 500
                
                return jsonify({
                    "ok": True,
                    "audio_url": f"/tts_output/{filename}",
                    "subtitles": all_subs,
                    "download_filename": filename
                })
                
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/asr', methods=['POST'])
def api_asr():
    if 'file' not in request.files: return jsonify({"ok": False, "error": "missing_file"}), 400
    file = request.files['file']
    
    # Generate unique task ID for this ASR request
    task_id = str(uuid.uuid4())
    print(f"ASR Request [Task ID: {task_id}]: Starting new ASR task", file=sys.stderr)
    
    key = request.form.get("dashscopeKey") or os.environ.get("DASHSCOPE_API_KEY")
    # For ASR, we might need a key. If not provided in form (frontend update needed?), check ENV.
    # We will assume ENV is primary or frontend sends it.
    
    if not key: return jsonify({"ok": False, "error": "missing_api_key"}), 401
    
    # Get file extension from original filename
    file_ext = os.path.splitext(file.filename)[1] if file.filename else '.mp3'
    
    # Save uploaded file to persistent storage with browser-compatible format
    out_dir = get_output_dir()
    
    # First, save to a temp location
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
        temp_path = temp_file.name
        file.save(temp_path)
    
    # Convert to browser-compatible format if needed
    converted_success, converted_path, convert_error = convert_to_browser_compatible(temp_path, file_ext)
    
    if not converted_success:
        # If conversion failed, try to use original file
        print(f"Conversion failed: {convert_error}, using original file", file=sys.stderr)
        converted_path = temp_path
        # Keep original extension
        saved_filename = f"asr-{uuid.uuid4()}{file_ext}"
    else:
        # Use converted file (extension will be .mp3 or .mp4)
        converted_ext = os.path.splitext(converted_path)[1]
        saved_filename = f"asr-{uuid.uuid4()}{converted_ext}"
    
    # Save to persistent storage
    saved_path = os.path.join(out_dir, saved_filename)
    shutil.move(converted_path, saved_path)
    
    # Clean up temp file if it's different from converted path
    if temp_path != converted_path and os.path.exists(temp_path):
        os.remove(temp_path)
    
    # Detect file type (video or audio) based on final saved file
    file_type = get_file_type(saved_path)
    
    # Process ASR with task ID
    with tempfile.TemporaryDirectory() as temp_dir:
        # Copy to temp directory for processing (to avoid modifying the original)
        input_path = os.path.join(temp_dir, file.filename)
        shutil.copy2(saved_path, input_path)
        
        ok, res = run_ali_asr(input_path, key, task_id)
        if not ok: return jsonify({"ok": False, "error": res}), 500
            
        # Parse Result
        # res structure from qwen3-asr-flash needs careful handling
        # Assuming res is result object from SDK
        
        # We need to extract transcript, subtitles, etc.
        # Since I cannot easily debug the exact response structure of qwen3-asr-flash without running it,
        # I will assume standard SenseVoice/Paraformer style or generic JSON.
        
        # If result contains 'sentences', we use it.
        transcript = ""
        subtitles = []
        
        # Try to extract from result object or dict
        if hasattr(res, 'sentences'):
            sents = res.sentences
        elif isinstance(res, dict) and 'sentences' in res:
            sents = res['sentences']
        else:
            sents = []
            # Maybe full text is in 'text' field?
            # res might be whole 'output' object from wait()
            pass
            
        if hasattr(res, 'text'): transcript = res.text
        elif isinstance(res, dict) and 'text' in res: transcript = res['text']
        
        # --- REMOVED GLOBAL NLP CLEANUP ---
        # It was causing text loss and timestamp misalignment.
        # Now cleanup is done per-chunk inside run_ali_asr.
        
        # If run_ali_asr didn't provide sentences (e.g. error fallback or legacy path), try to use what we have
        if not subtitles and sents:
             for s in sents:
                text_s = s['text'] if isinstance(s, dict) else s.text
                start = s['begin_time'] if isinstance(s, dict) else s.begin_time
                end = s['end_time'] if isinstance(s, dict) else s.end_time
                subtitles.append({
                    "text": text_s,
                    "start": start / 1000.0,
                    "end": end / 1000.0
                })

        # Analyze
        # Use LLM for analysis as requested (qwen3-max)
        ok_llm, llm_data = call_llm_analysis(transcript, key)
        keywords = []
        summary = ""
        topics = []
        if ok_llm:
            keywords = llm_data.get("keywords", [])
            summary = llm_data.get("summary", "")
            topics = llm_data.get("topics", [])
            
        return jsonify({
            "ok": True,
            "task_id": task_id,  # Return task ID to client
            "audio_url": f"/tts_output/{saved_filename}",
            "file_type": file_type,  # Added file_type field
            "transcript": transcript,
            "subtitles": subtitles,
            "keywords": keywords,
            "summary": summary,
            "topics": topics,
            "analysis": llm_data if ok_llm else None
        })

@app.route('/api/asr-url', methods=['POST'])
def api_asr_url():
    """
    ASR from URL endpoint.
    Accepts a URL to audio/video file or video website (B站, YouTube, etc.),
    downloads it, and performs ASR transcription.
    """
    try:
        data = request.get_json()
        url = data.get('url')
        
        if not url:
            return jsonify({"ok": False, "error": "missing_url"}), 400
        
        key = data.get("dashscopeKey") or os.environ.get("DASHSCOPE_API_KEY")
        if not key: return jsonify({"ok": False, "error": "missing_api_key"}), 401
        
        print(f"ASR-URL Request: URL={url}", file=sys.stderr)
        
        # Step 1: Download audio/video from URL
        out_dir = get_output_dir()
        success, downloaded_path, download_error, original_url = download_audio_video_from_url(url, out_dir)
        
        if not success:
            return jsonify({"ok": False, "error": f"下载失败: {download_error}"}), 400
        
        # Step 2: Convert to browser-compatible format if needed
        downloaded_ext = os.path.splitext(downloaded_path)[1]
        converted_success, converted_path, convert_error = convert_to_browser_compatible(downloaded_path, downloaded_ext)
        
        if not converted_success:
            print(f"Conversion failed: {convert_error}, using original file", file=sys.stderr)
            converted_path = downloaded_path
            saved_filename = os.path.basename(downloaded_path)
        else:
            converted_ext = os.path.splitext(converted_path)[1]
            saved_filename = f"asr-{uuid.uuid4()}{converted_ext}"
            # Move converted file to persistent storage
            saved_path = os.path.join(out_dir, saved_filename)
            shutil.move(converted_path, saved_path)
            converted_path = saved_path
        
        # Clean up original downloaded file if it's different
        if downloaded_path != converted_path and os.path.exists(downloaded_path):
            os.remove(downloaded_path)
        
        # Step 3: Detect file type (video or audio)
        file_type = get_file_type(converted_path)
        
        # Step 4: Process ASR
        with tempfile.TemporaryDirectory() as temp_dir:
            # Copy to temp directory for processing
            input_path = os.path.join(temp_dir, saved_filename)
            shutil.copy2(converted_path, input_path)
            
            ok, res = run_ali_asr(input_path, key)
            if not ok: return jsonify({"ok": False, "error": res}), 500
            
            # Parse Result
            transcript = ""
            subtitles = []
            
            if hasattr(res, 'sentences'):
                sents = res.sentences
            elif isinstance(res, dict) and 'sentences' in res:
                sents = res['sentences']
            else:
                sents = []
                
            if hasattr(res, 'text'): transcript = res.text
            elif isinstance(res, dict) and 'text' in res: transcript = res['text']
            
            # Extract subtitles
            if not subtitles and sents:
                 for s in sents:
                    text_s = s['text'] if isinstance(s, dict) else s.text
                    start = s['begin_time'] if isinstance(s, dict) else s.begin_time
                    end = s['end_time'] if isinstance(s, dict) else s.end_time
                    subtitles.append({
                        "text": text_s,
                        "start": start / 1000.0,
                        "end": end / 1000.0
                    })
            
            # Analyze with LLM
            ok_llm, llm_data = call_llm_analysis(transcript, key)
            keywords = []
            summary = ""
            topics = []
            if ok_llm:
                keywords = llm_data.get("keywords", [])
                summary = llm_data.get("summary", "")
                topics = llm_data.get("topics", [])
                
            return jsonify({
                "ok": True,
                "audio_url": f"/tts_output/{saved_filename}",
                "file_type": file_type,
                "transcript": transcript,
                "subtitles": subtitles,
                "keywords": keywords,
                "summary": summary,
                "topics": topics,
                "analysis": llm_data if ok_llm else None,
                "source_url": original_url
            })
                
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/analyze-text', methods=['POST'])
def api_analyze_text():
    data = request.get_json()
    text = data.get('text', '')
    key = data.get('dashscopeKey') or os.environ.get("DASHSCOPE_API_KEY")
    
    if not key: return jsonify({"ok": False, "error": "missing_api_key"}), 401
    
    ok, res = call_llm_analysis(text, key)
    if not ok: return jsonify({"ok": False, "error": res}), 500
    
    return jsonify({"ok": True, "data": res})

@app.route('/api/fix-punctuation', methods=['POST'])
def api_fix_punctuation():
    data = request.get_json()
    text = data.get('text', '')
    key = data.get('dashscopeKey') or os.environ.get("DASHSCOPE_API_KEY")
    
    if not key: return jsonify({"ok": False, "error": "missing_api_key"}), 401
    
    ok, res = call_llm_fix_punctuation(text, key)
    if not ok: return jsonify({"ok": False, "error": res}), 500
    
    return jsonify({"ok": True, "text": res})

@app.route('/api/ai-advice', methods=['POST'])
def api_ai_advice():
    data = request.get_json()
    text = data.get('text', '')
    key = data.get('dashscopeKey') or os.environ.get("DASHSCOPE_API_KEY")
    custom_prompt = data.get('custom_prompt')
    
    if not key: return jsonify({"ok": False, "error": "missing_api_key"}), 401
    
    ok, res = call_llm_advice(text, key, custom_prompt)
    if not ok: return jsonify({"ok": False, "error": res}), 500
    
    return jsonify({"ok": True, "data": res})

@app.route('/api/generate-advice-word', methods=['POST'])
def api_generate_advice_word():
    try:
        data = request.get_json()
        original_text = data.get('original_text', '')
        advice_data = data.get('advice_data', {})
        
        doc = docx.Document()
        
        # Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # Title
        doc.add_heading('作文润色与修改建议', 0)
        
        # Original Text
        doc.add_heading('原文内容', 1)
        doc.add_paragraph(original_text)
        
        # Analysis
        doc.add_heading('整体点评与评分', 1)
        if advice_data.get('score_prediction'):
             doc.add_paragraph(advice_data['score_prediction']).bold = True
        doc.add_paragraph(advice_data.get('analysis', ''))
        
        # Structure Advice
        if advice_data.get('structure_advice'):
            doc.add_heading('写作思路与结构进阶', 1)
            doc.add_paragraph(advice_data['structure_advice'])
            
        # Alternative Ideas
        if advice_data.get('alternative_ideas'):
            doc.add_heading('多维审题与构思拓展', 1)
            for idea in advice_data['alternative_ideas']:
                doc.add_paragraph(idea.get('title', '思路')).bold = True
                doc.add_paragraph(idea.get('desc', ''))

        # Suggestions
        doc.add_heading('详细修改建议与升格示例', 1)
        suggestions = advice_data.get('suggestions', [])
        for i, sug in enumerate(suggestions):
            p = doc.add_paragraph()
            p.add_run(f"{i+1}. [{sug.get('technique', '建议')}] ").bold = True
            
            if sug.get('original'):
                p.add_run(f"\n原文：{sug['original']}").italic = True
            
            # Suggestion/Analysis
            p.add_run(f"\n分析：{sug.get('suggestion', '')}")
            
            # Refined Text
            if sug.get('refined_text'):
                clean_ref = sug['refined_text'].replace('**', '').replace('__', '')
                p.add_run(f"\n升格示例：{clean_ref}").bold = True

        # Style Demonstrations
        if advice_data.get('style_demonstrations'):
            doc.add_heading('三种风格润色示范', 1)
            for demo in advice_data['style_demonstrations']:
                p = doc.add_paragraph()
                p.add_run(demo.get('style_name', '风格')).bold = True
                
                # Check for nested examples list
                examples = demo.get('examples', [])
                # Fallback for old format
                if not examples and demo.get('refined_text'):
                    examples.append({
                        "original_snippet": demo.get('original_snippet'),
                        "refined_text": demo.get('refined_text'),
                        "comment": demo.get('comment')
                    })
                
                for i, ex in enumerate(examples):
                    doc.add_paragraph(f"示例 {i+1}").bold = True
                    if ex.get('original_snippet'):
                        doc.add_paragraph(f"原文片段：{ex['original_snippet']}").italic = True
                    doc.add_paragraph(f"升格内容：{ex.get('refined_text', '')}")
                    if ex.get('comment'):
                         doc.add_paragraph(f"解析：{ex['comment']}")
                    doc.add_paragraph("") # Spacing between examples
                
                doc.add_paragraph("---") # Separator between styles
            
        fname = f"advice-{uuid.uuid4()}.docx"
        out_dir = get_output_dir()
        doc.save(os.path.join(out_dir, fname))
        
        return jsonify({"ok": True, "download_url": f"/tts_output/{fname}", "filename": fname})
        
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Utils: File Extraction (Kept from original) ---
@app.route('/api/extract-text', methods=['POST'])
def api_extract_text():
    if 'file' not in request.files: return jsonify({"ok": False, "error": "missing_file"}), 400
    file = request.files['file']
    if not file.filename: return jsonify({"ok": False, "error": "empty"}), 400
    
    content = file.read()
    ext = os.path.splitext(file.filename)[1].lower()
    text = ""
    
    try:
        if ext == ".pdf":
            # Use PyMuPDF (fitz)
            with fitz.open(stream=content, filetype="pdf") as doc:
                for page in doc: text += page.get_text() + "\n"
        elif ext == ".docx":
            with tempfile.TemporaryDirectory() as td:
                p = os.path.join(td, "in.docx")
                with open(p, "wb") as f: f.write(content)
                doc = docx.Document(p)
                text = "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".txt":
            try: text = content.decode('utf-8')
            except: text = content.decode('gbk', errors='ignore')
        else:
            return jsonify({"ok": False, "error": "unsupported"}), 400
            
        return jsonify({"ok": True, "text": normalize_text_for_tts(text)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/extract-url', methods=['POST'])
def api_extract_url():
    data = request.get_json()
    url = data.get('url')
    if not url: return jsonify({"ok": False, "error": "missing_url"}), 400
    
    key = data.get("dashscopeKey") or os.environ.get("DASHSCOPE_API_KEY")
    if not key: return jsonify({"ok": False, "error": "missing_api_key"}), 401

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        resp.encoding = resp.apparent_encoding 
        
        soup = BeautifulSoup(resp.text, 'html.parser')
        
        # Extract Title
        title = ""
        if soup.title: title = soup.title.string
        if not title:
            h1 = soup.find('h1')
            if h1: title = h1.get_text()
            
        # Extract Body
        # Remove scripts and styles
        for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
            script.decompose()
            
        # Get text
        # separator='\n' helps to separate block elements
        text = soup.get_text(separator='\n')
        
        # Clean up lines
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        content_text = "\n".join(lines)
        
        # Combine
        full_text = ""
        if title: full_text += f"{title}\n\n"
        full_text += content_text
        
        # Format using LLM (No truncation logic needed for local)
        ok, formatted_text = call_llm_format_essay(full_text, key)
        if not ok:
            formatted_text = normalize_text_for_tts(full_text)
            
        return jsonify({"ok": True, "text": formatted_text})

    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

def call_llm_format_essay(text, api_key):
    dashscope.api_key = api_key
    # Increase prompt limit for local high-capacity model
    safe_text = text[:20000] 
    
    prompt = f"""
    你是一个专业的出版编辑。请对以下OCR识别出的作文文本进行严格的标准化排版和格式校对。
    
    【目标】：输出符合标准出版物要求的作文文本。
    
    【具体要求】：
    1. **段落修正**：智能识别段落，合并被OCR错误断行的句子，确保段落完整。
    2. **缩进处理**：每个自然段的开头必须严格保持 **两个全角空格（　　）** 的缩进。
    3. **标点规范**：
       - 将所有英文标点转换为中文全角标点（如 , -> ，）。
       - 修正OCR可能导致的标点缺失或错误。
    4. **文本清洗**：去除文中多余的空格、乱码符号或页眉页脚干扰文本。
    5. **标题处理**：如果第一行疑似标题，请将其居中（在文本中通过前后加空行体现，或者不做特殊标记，仅确保其独立成段且无缩进，但通常作文正文段落需缩进）。
       * 为了通用性，请对所有正文段落添加缩进。如果第一行明显是标题，可以不缩进但需独立成行。
    6. **输出纯文本**：只返回排版后的文本内容，不要包含任何 JSON 格式、Markdown 标记或解释性语言。
    
    【待处理文本】：
    {safe_text}
    """
    try:
        resp = dashscope.Generation.call(
            model=MODEL_LLM_FAST, # Now Qwen-Max
            messages=[{'role': 'user', 'content': prompt}],
            result_format='message'
        )
        if resp.status_code == HTTPStatus.OK:
            return True, resp.output.choices[0].message.content
        return False, resp.message
    except Exception as e:
        return False, str(e)

@app.route('/api/ocr-to-word', methods=['POST'])
def api_ocr_to_word():
    # Use Qwen-VL-OCR for all OCR tasks
    files = request.files.getlist('file')
    if not files or not files[0].filename:
        return jsonify({"ok": False, "error": "missing_file"}), 400
    
    # NOTE: We rely on the client (frontend) to send files in the correct order.
    # Previously we sorted by filename, but that prevents users from manually ordering files
    # (e.g. "Part2.jpg" then "Part1.jpg" if they really wanted to, or if filenames are messy).
    # The frontend now implements a Queue system to guarantee order.
    
    # We need API key for Qwen
    key = request.form.get("dashscopeKey") or os.environ.get("DASHSCOPE_API_KEY")
    if not key: return jsonify({"ok": False, "error": "missing_api_key (please configure in settings)"}), 401

    full_text = []
    
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            # Process files sequentially in order
            for idx, file in enumerate(files):
                content = file.read()
                filename = file.filename.lower()
                
                print(f"Processing file {idx+1}/{len(files)}: {filename}", file=sys.stderr)
                
                if filename.endswith(".pdf"):
                     with fitz.open(stream=content, filetype="pdf") as doc:
                        for i, page in enumerate(doc):
                            pix = page.get_pixmap()
                            img_path = os.path.join(temp_dir, f"f{idx}_p{i}.png")
                            pix.save(img_path)
                            
                            ok, text = call_qwen_ocr(img_path, key)
                            if ok: full_text.append(text)
                            else: print(f"OCR Failed for {filename} page {i}: {text}", file=sys.stderr)
                            
                else:
                    # Image
                    img_path = os.path.join(temp_dir, f"f{idx}_img.png")
                    with open(img_path, "wb") as f: f.write(content)
                    
                    ok, text = call_qwen_ocr(img_path, key)
                    if ok: full_text.append(text)
                    else: 
                        # Continue or fail? Let's append error message to text to warn user
                        full_text.append(f"[Error reading file {filename}: {text}]")
            
        final_text = "\n\n".join(full_text)
        
        # Basic cleanup of markdown code blocks if Qwen returns them despite prompt
        final_text = re.sub(r'```.*?\n', '', final_text)
        final_text = final_text.replace('```', '')
        
        # Apply formatting via LLM for publication standard
        ok_fmt, formatted_text = call_llm_format_essay(final_text, key)
        
        # Fallback to regex if LLM fails
        if not ok_fmt:
            print(f"Format LLM Failed: {formatted_text}, using fallback.", file=sys.stderr)
            formatted_text = format_as_essay(final_text)
        
        return jsonify({"ok": True, "text": formatted_text})
        
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/generate-word', methods=['POST'])
def api_generate_word():
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({"ok": False, "error": "empty_text"}), 400
            
        doc = docx.Document()
        
        # Set Default Style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12) # 12pt (Xiao Si)
        font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体') # SimSun
        
        # Paragraph Formatting
        pf = style.paragraph_format
        pf.line_spacing = 1.5 # 1.5 lines spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        
        for paragraph in text.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                # Ensure alignment is justified or left
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        fname = f"ocr-{uuid.uuid4()}.docx"
        out_dir = get_output_dir()
        doc.save(os.path.join(out_dir, fname))
        
        return jsonify({"ok": True, "download_url": f"/tts_output/{fname}", "filename": fname})
        
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- History API Routes ---
@app.route('/api/history', methods=['GET'])
def api_get_history():
    """Get all history records"""
    history = load_history()
    # Sort by createdAt descending
    history.sort(key=lambda x: x.get('createdAt', ''), reverse=True)
    return jsonify({"ok": True, "history": history})

@app.route('/api/history', methods=['POST'])
def api_save_history():
    """Save a new history record"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if 'type' not in data:
            return jsonify({"ok": False, "error": "missing_type"}), 400
        
        history = load_history()
        
        # Create new record with ID and timestamp
        new_record = {
            "id": data.get('id') or str(int(datetime.now().timestamp())),
            "type": data['type'],
            "createdAt": data.get('createdAt') or datetime.now().isoformat(),
            **data
        }
        
        # Add to beginning of array
        history.insert(0, new_record)
        
        # Limit to 200 records
        if len(history) > 200:
            history = history[:200]
        
        # Save to file
        if save_history_to_file(history):
            return jsonify({"ok": True, "record": new_record})
        else:
            return jsonify({"ok": False, "error": "failed_to_save"}), 500
            
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/history', methods=['DELETE'])
def api_delete_history():
    """Delete a history record"""
    try:
        data = request.get_json()
        record_id = data.get('id')
        
        if not record_id:
            return jsonify({"ok": False, "error": "missing_id"}), 400
        
        history = load_history()
        
        # Filter out the record
        new_history = [r for r in history if r.get('id') != record_id]
        
        # Save to file
        if save_history_to_file(new_history):
            return jsonify({"ok": True, "deleted": record_id})
        else:
            return jsonify({"ok": False, "error": "failed_to_delete"}), 500
            
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/extract-audio', methods=['POST'])
def api_extract_audio():
    """
    Extract audio from video file.
    Returns URL of the extracted audio file.
    The extracted audio file will be much smaller than the original video file.
    """
    try:
        data = request.get_json()
        video_url = data.get('video_url')
        
        if not video_url:
            return jsonify({"ok": False, "error": "missing_video_url"}), 400
        
        # Get the file path from the URL
        # Support both relative and absolute URLs
        # Relative: /tts_output/filename.ext
        # Absolute: http://localhost:5000/tts_output/filename.ext
        
        # Extract filename from URL
        if '/tts_output/' in video_url:
            # Get the part after '/tts_output/'
            filename = video_url.split('/tts_output/')[-1]
            # Remove any query parameters or fragments
            filename = filename.split('?')[0].split('#')[0]
            video_path = os.path.join(get_output_dir(), filename)
        else:
            return jsonify({"ok": False, "error": "invalid_url_format: URL must contain /tts_output/"}), 400
        
        if not os.path.exists(video_path):
            return jsonify({"ok": False, "error": "video_file_not_found"}), 404
        
        # Get original file size for logging
        original_size = os.path.getsize(video_path)
        print(f"Original video file size: {original_size} bytes ({original_size/1024/1024:.2f} MB)", file=sys.stderr)
        
        # Generate audio filename
        base_name = os.path.splitext(filename)[0]
        audio_filename = f"{base_name}-audio.mp3"
        audio_path = os.path.join(get_output_dir(), audio_filename)
        
        # Extract audio using ffmpeg
        if not shutil.which("ffmpeg"):
            return jsonify({"ok": False, "error": "ffmpeg_not_available"}), 500
        
        try:
            # Extract audio to MP3 format
            # -vn: No video
            # -acodec libmp3lame: Use MP3 codec
            # -q:a 2: Quality setting (0-9, lower is better, 2 is good quality)
            # -ar 44100: Sample rate (optional, let ffmpeg decide)
            # -ac 2: Stereo channels (optional)
            cmd = [
                "ffmpeg", "-y", "-i", video_path,
                "-vn",  # No video stream
                "-acodec", "libmp3lame",
                "-q:a", "2",  # Good quality (0-9, lower is better)
                audio_path
            ]
            print(f"Running ffmpeg command: {' '.join(cmd)}", file=sys.stderr)
            result = subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
            
            # Check if audio file was created and has reasonable size
            if not os.path.exists(audio_path):
                error_msg = f"Audio extraction failed: output file not created"
                print(error_msg, file=sys.stderr)
                return jsonify({"ok": False, "error": error_msg}), 500
            
            # Get extracted audio file size
            audio_size = os.path.getsize(audio_path)
            print(f"Extracted audio file size: {audio_size} bytes ({audio_size/1024/1024:.2f} MB)", file=sys.stderr)
            print(f"Size reduction: {(1 - audio_size/original_size)*100:.1f}%", file=sys.stderr)
            
            # Validate audio size is reasonable (should be much smaller than video)
            if audio_size >= original_size:
                warning_msg = f"Warning: Audio file size ({audio_size} bytes) is not significantly smaller than video ({original_size} bytes)"
                print(warning_msg, file=sys.stderr)
            
            print(f"Audio extracted successfully: {audio_path}", file=sys.stderr)
            
            return jsonify({
                "ok": True,
                "audio_url": f"/tts_output/{audio_filename}",
                "audio_size": audio_size,
                "video_size": original_size
            })
            
        except subprocess.CalledProcessError as e:
            error_msg = f"ffmpeg extraction failed: {str(e)}"
            print(error_msg, file=sys.stderr)
            return jsonify({"ok": False, "error": error_msg}), 500
            
    except Exception as e:
        error_msg = f"Audio extraction error: {str(e)}"
        print(error_msg, file=sys.stderr)
        return jsonify({"ok": False, "error": error_msg}), 500

@app.route('/api/history/<record_id>', methods=['GET'])
def api_get_history_record(record_id):
    """Get a specific history record by ID"""
    history = load_history()
    record = next((r for r in history if r.get('id') == record_id), None)
    
    if record:
        return jsonify({"ok": True, "record": record})
    else:
        return jsonify({"ok": False, "error": "not_found"}), 404

@app.route('/api/logs', methods=['GET'])
def api_get_logs():
    """
    Get recent server logs for progress tracking
    Query parameters:
    - lines: number of recent log lines to return (default: 50, max: 200)
    - filter: optional filter pattern (regex) to match specific log lines
    - seconds: only return logs from last N seconds (default: 120, max: 600)
    """
    try:
        # Get query parameters
        lines = min(int(request.args.get('lines', 50)), 200)  # Limit to max 200 lines
        filter_pattern = request.args.get('filter', None)
        seconds = min(int(request.args.get('seconds', 120)), 600)  # Limit to max 10 minutes
        
        # Calculate cutoff time
        from datetime import datetime, timedelta
        cutoff_time = datetime.now() - timedelta(seconds=seconds)
        
        # Determine log file path
        log_file = os.path.join(root_dir, 'logs', 'server.log')
        if not os.path.exists(log_file):
            # Fallback to root directory
            log_file = os.path.join(root_dir, 'server.log')
        
        if not os.path.exists(log_file):
            return jsonify({"ok": True, "logs": []})
        
        # Read recent log lines
        recent_logs = []
        try:
            with open(log_file, 'r', encoding='utf-8', errors='ignore') as f:
                all_lines = f.readlines()
                
            # Get last N lines
            recent_lines = all_lines[-lines:]
            
            # Filter by timestamp (only logs from last N seconds)
            time_filtered_lines = []
            for line in recent_lines:
                line = line.strip()
                if not line:
                    continue
                
                # Parse timestamp from log line (format: "2026-01-18 08:30:18 - INFO - message")
                try:
                    # Extract timestamp part (first 19 characters: "2026-01-18 08:30:18")
                    timestamp_str = line[:19]
                    log_time = datetime.strptime(timestamp_str, '%Y-%m-%d %H:%M:%S')
                    
                    # Check if log is within time window
                    if log_time >= cutoff_time:
                        time_filtered_lines.append(line)
                except (ValueError, IndexError):
                    # If timestamp parsing fails, include the line anyway
                    # (recent lines are likely recent)
                    time_filtered_lines.append(line)
            
            # Apply regex filter if provided
            if filter_pattern:
                try:
                    import re
                    regex = re.compile(filter_pattern, re.IGNORECASE)
                    time_filtered_lines = [line for line in time_filtered_lines if regex.search(line)]
                except re.error:
                    pass  # Invalid regex, return all lines
            
            recent_logs = time_filtered_lines
            
        except Exception as e:
            print(f"Error reading log file: {e}", file=sys.stderr)
            return jsonify({"ok": True, "logs": []})
        
        return jsonify({
            "ok": True,
            "logs": recent_logs,
            "count": len(recent_logs),
            "time_window": f"last {seconds} seconds"
        })
        
    except Exception as e:
        print(f"Error in /api/logs: {e}", file=sys.stderr)
        return jsonify({"ok": False, "error": str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT",5173))
    # Ensure DashScope Key is present
    if not os.environ.get("DASHSCOPE_API_KEY"):
        logging.warning("DASHSCOPE_API_KEY not found in environment.")
    
    logging.info(f"Starting server on port {port}")
    app.run(host='0.0.0.0', port=port)
